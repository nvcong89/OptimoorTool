"""Generate mooring analysis Word reports from Optimoor input/output Excel files."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

import openpyxl
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Sheet names in OptimoorTool_Input.xlsm
SH_BERTH = "1A.Berth_Data"
SH_VESSEL = "2A.Vessel_Data"
SH_ENV = "3A.GeneralEnvir."
SH_RUN_CASE = "3B.RunCase"

# Sheet names in OutputResult.xlsx (from optimoor_rtf_to_excel.py)
SH_CASES = "Cases"
SH_GREATEST_LINE = "Greatest_Line"
SH_GREATEST_BOLLARD = "Greatest_Bollard"
SH_HOOK_BOLLARD = "Hook_Bollard_Forces"
SH_LINE_TENSIONS = "Line_Tensions"

MBL_LIMIT = 0.50  # OCIMF 50% MBL allowable limit

REQUIRED_INPUT_SHEETS = [SH_BERTH, SH_VESSEL, SH_ENV, SH_RUN_CASE]
REQUIRED_OUTPUT_SHEETS = [
    SH_CASES,
    SH_GREATEST_LINE,
    SH_GREATEST_BOLLARD,
    SH_HOOK_BOLLARD,
    SH_LINE_TENSIONS,
]


@dataclass
class ReportMetadata:
    """Optional cover-page metadata supplied by the user."""

    project_name: str = "Gemalink Container Terminal Phase 2"
    reference: str = ""
    report_date: str = ""
    client: str = "Gemalink International Terminal"
    author: str = "OptimoorTool"


@dataclass
class InputReportData:
    """Parsed data from OptimoorTool_Input.xlsm."""

    berth_rows: list[dict[str, str]] = field(default_factory=list)
    vessel_rows: list[dict[str, str]] = field(default_factory=list)
    environment_rows: list[dict[str, str]] = field(default_factory=list)
    scenarios: list[dict[str, Any]] = field(default_factory=list)
    vessel_name: str = ""
    vessel_id: str = ""
    mbl_ton: float | None = None


@dataclass
class OutputReportData:
    """Parsed data from OutputResult.xlsx."""

    cases: pd.DataFrame = field(default_factory=pd.DataFrame)
    greatest_lines: pd.DataFrame = field(default_factory=pd.DataFrame)
    greatest_bollards: pd.DataFrame = field(default_factory=pd.DataFrame)
    hook_bollard: pd.DataFrame = field(default_factory=pd.DataFrame)
    line_tensions: pd.DataFrame = field(default_factory=pd.DataFrame)
    bollard_matrix: list[dict[str, Any]] = field(default_factory=list)
    case_count: int = 0


def _clean(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    if text in {"", "None", "#N/A", "nan"}:
        return ""
    return text


def _format_num(value: Any, decimals: int = 2) -> str:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    try:
        num = float(value)
        if num == int(num) and decimals <= 0:
            return str(int(num))
        return f"{num:.{decimals}f}".rstrip("0").rstrip(".")
    except (TypeError, ValueError):
        return _clean(value)


def _format_pct(value: Any) -> str:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    try:
        num = float(value)
        if num <= 1:
            return f"{num * 100:.1f}%"
        return f"{num:.1f}%"
    except (TypeError, ValueError):
        return _clean(value)


def validate_input_files(input_path: Path, output_path: Path) -> list[str]:
    """Return a list of validation errors; empty list means OK."""
    errors: list[str] = []

    if not input_path.exists():
        errors.append(f"Input file not found: {input_path}")
    if not output_path.exists():
        errors.append(f"Output file not found: {output_path}")

    if input_path.exists():
        try:
            wb = openpyxl.load_workbook(input_path, read_only=True, keep_vba=True)
            missing = [s for s in REQUIRED_INPUT_SHEETS if s not in wb.sheetnames]
            if missing:
                errors.append(f"Input workbook missing sheets: {', '.join(missing)}")
            wb.close()
        except Exception as exc:
            errors.append(f"Cannot read input workbook: {exc}")

    if output_path.exists():
        try:
            xls = pd.ExcelFile(output_path)
            missing = [s for s in REQUIRED_OUTPUT_SHEETS if s not in xls.sheet_names]
            if missing:
                errors.append(f"Output workbook missing sheets: {', '.join(missing)}")
        except Exception as exc:
            errors.append(f"Cannot read output workbook: {exc}")

    return errors


def parse_berth_data(wb: openpyxl.Workbook) -> list[dict[str, str]]:
    """Extract berth parameters from 1A.Berth_Data."""
    ws = wb[SH_BERTH]
    rows: list[dict[str, str]] = []
    skip_headers = {"BERTH PARAMETERS", "FENDER SYSTEM", "BOLLARD SYSTEM", "Unit", "Value"}

    for row in ws.iter_rows(min_row=3, max_col=3, values_only=True):
        label, unit, value = (row + (None, None, None))[:3]
        label_s = _clean(label)
        if not label_s or label_s in skip_headers:
            continue
        if _clean(unit) == "Unit" and _clean(value) == "Value":
            continue
        rows.append({
            "parameter": label_s,
            "unit": _clean(unit) if _clean(unit) != "-" else "",
            "value": _format_num(value) if isinstance(value, (int, float)) else _clean(value),
        })
    return rows


def _find_vessel_column(ws, vessel_id: str | None = None) -> tuple[int, str]:
    """Return (column_index, vessel_id) for the primary vessel column (1-based col C=3)."""
    header_row = list(ws.iter_rows(min_row=3, max_row=3, values_only=True))[0]
    for idx, header in enumerate(header_row):
        if idx < 2:
            continue
        name = _clean(header)
        if not name:
            continue
        if vessel_id and name.lower() != vessel_id.lower():
            continue
        return idx, name
    return 2, _clean(header_row[2]) if len(header_row) > 2 else "Vessel 1"


def parse_vessel_data(wb: openpyxl.Workbook, vessel_id: str | None = None) -> tuple[list[dict[str, str]], str, str, float | None]:
    """Extract vessel parameters from 2A.Vessel_Data."""
    ws = wb[SH_VESSEL]
    col_idx, vid = _find_vessel_column(ws, vessel_id)
    rows: list[dict[str, str]] = []
    mbl_ton: float | None = None
    skip = {"Parameters", "Unit", "VESSEL DATA", "MOORING LINE CHARACTERISTIC"}

    for row in ws.iter_rows(min_row=4, max_col=col_idx + 1, values_only=True):
        label = _clean(row[0]) if row else ""
        unit = _clean(row[1]) if row and len(row) > 1 else ""
        value = row[col_idx] if row and len(row) > col_idx else None
        if not label or label in skip or label.startswith("Note for"):
            continue
        if label.startswith(("Tail là", "Nếu không", "Lấy Winch", "Brake limit", "Theo OCIMF", "Không có")):
            break

        val_str = _format_num(value) if isinstance(value, (int, float)) else _clean(value)
        rows.append({"parameter": label, "unit": unit, "value": val_str})

        if "Minimum breaking load of mooring line" in label and isinstance(value, (int, float)):
            mbl_ton = float(value)

    vessel_name = vid
    return rows, vessel_name, vid, mbl_ton


def parse_environment(wb: openpyxl.Workbook) -> list[dict[str, str]]:
    """Extract environmental conditions from 3A.GeneralEnvir."""
    ws = wb[SH_ENV]
    rows: list[dict[str, str]] = []

    for row in ws.iter_rows(min_row=6, max_row=14, values_only=True):
        cells = list(row)
        if not cells:
            continue

        # Wind block (col B-F)
        if _clean(cells[1]) == "Condition" and "Wind" in _clean(cells[2]):
            continue
        if len(cells) > 5 and _clean(cells[1]) == "Operating":
            rows.append({
                "parameter": "Wind speed (1-min, operating)",
                "unit": "knot",
                "value": _format_num(cells[5], 2),
            })

        # Current block (row label in col B)
        if len(cells) > 5 and _clean(cells[1]) == "Operating" and isinstance(cells[2], (int, float)):
            rows.append({
                "parameter": "Current velocity (operating)",
                "unit": "knot",
                "value": _format_num(cells[5], 2),
            })
            rows.append({
                "parameter": "Current direction to True North",
                "unit": "deg",
                "value": _format_num(cells[4], 0),
            })

        # Wave block (Hs in col J when col I is Operating)
        if len(cells) > 9 and _clean(cells[8]) == "Operating":
            rows.append({
                "parameter": "Significant wave height Hs (operating)",
                "unit": "m",
                "value": _format_num(cells[9], 2),
            })

    return rows


def parse_scenarios(wb: openpyxl.Workbook) -> list[dict[str, Any]]:
    """Extract design scenarios from 3B.RunCase."""
    ws = wb[SH_RUN_CASE]
    scenarios: list[dict[str, Any]] = []

    for row in ws.iter_rows(min_row=5, values_only=True):
        name = _clean(row[0]) if row else ""
        if not name or not re.match(r".*case\s*\d+", name, re.I):
            continue

        case_no = row[1] if len(row) > 1 else None
        vessel_id = _clean(row[3]) if len(row) > 3 else ""
        loading = _clean(row[4]) if len(row) > 4 else ""
        wind_dir = _clean(row[6]) if len(row) > 6 else ""
        current_dir = _clean(row[7]) if len(row) > 7 else ""
        tide_label = _clean(row[8]) if len(row) > 8 else ""
        draft = row[10] if len(row) > 10 else ""
        wind_knots = row[13] if len(row) > 13 else ""
        current_knots = row[14] if len(row) > 14 else ""
        water_level = row[16] if len(row) > 16 else ""

        scenarios.append({
            "no": int(case_no) if case_no is not None else len(scenarios) + 1,
            "name": name,
            "vessel": vessel_id,
            "loading": loading,
            "draft": _format_num(draft, 1),
            "water_level": _format_num(water_level, 2) if isinstance(water_level, (int, float)) else tide_label,
            "tide_label": tide_label,
            "wind_direction": wind_dir,
            "current_direction": current_dir,
            "wind_speed": _format_num(wind_knots, 2),
            "current_speed": _format_num(current_knots, 2),
        })

    return scenarios


def read_input_workbook(path: Path) -> InputReportData:
    """Load and parse OptimoorTool_Input.xlsm."""
    wb = openpyxl.load_workbook(path, read_only=True, keep_vba=True, data_only=True)
    try:
        scenarios = parse_scenarios(wb)
        vessel_id = scenarios[0]["vessel"] if scenarios else None
        vessel_rows, vessel_name, vid, mbl = parse_vessel_data(wb, vessel_id)
        return InputReportData(
            berth_rows=parse_berth_data(wb),
            vessel_rows=vessel_rows,
            environment_rows=parse_environment(wb),
            scenarios=scenarios,
            vessel_name=vessel_name,
            vessel_id=vid,
            mbl_ton=mbl,
        )
    finally:
        wb.close()


def read_output_workbook(path: Path) -> OutputReportData:
    """Load and parse OutputResult.xlsx."""
    cases = pd.read_excel(path, sheet_name=SH_CASES)
    greatest_lines = pd.read_excel(path, sheet_name=SH_GREATEST_LINE)
    greatest_bollards = pd.read_excel(path, sheet_name=SH_GREATEST_BOLLARD)
    hook_bollard = pd.read_excel(path, sheet_name=SH_HOOK_BOLLARD)
    line_tensions = pd.read_excel(path, sheet_name=SH_LINE_TENSIONS)

    case_count = int(cases["Batch"].nunique()) if not cases.empty else 0
    bollard_matrix = build_bollard_matrix(hook_bollard, case_count)

    return OutputReportData(
        cases=cases,
        greatest_lines=greatest_lines,
        greatest_bollards=greatest_bollards,
        hook_bollard=hook_bollard,
        line_tensions=line_tensions,
        bollard_matrix=bollard_matrix,
        case_count=case_count,
    )


def build_bollard_matrix(hook_bollard: pd.DataFrame, case_count: int) -> list[dict[str, Any]]:
    """Pivot Hook/Bollard forces into a point × case matrix (tonnes)."""
    if hook_bollard.empty or case_count == 0:
        return []

    df = hook_bollard.copy()
    df["Batch"] = pd.to_numeric(df["Batch"], errors="coerce")
    df["Total Force"] = pd.to_numeric(df["Total Force"], errors="coerce")

    # Keep bollards with at least one non-null force
    grouped = (
        df.groupby(["Hook/Bollard", "Batch"], as_index=False)["Total Force"]
        .max()
    )

    points = sorted(grouped["Hook/Bollard"].dropna().unique(), key=lambda x: (len(str(x)), str(x)))
    matrix: list[dict[str, Any]] = []

    for point in points:
        row: dict[str, Any] = {"point": str(point)}
        sub = grouped[grouped["Hook/Bollard"] == point]
        for case_no in range(1, case_count + 1):
            val = sub.loc[sub["Batch"] == case_no, "Total Force"]
            force = val.iloc[0] if not val.empty and pd.notna(val.iloc[0]) else None
            row[f"case_{case_no}"] = _format_num(force, 1) if force is not None else "-"
        matrix.append(row)

    return matrix


def build_greatest_line_rows(greatest_lines: pd.DataFrame) -> list[dict[str, str]]:
    """Format Greatest_Line sheet for the report table."""
    rows: list[dict[str, str]] = []
    if greatest_lines.empty:
        return rows

    for _, rec in greatest_lines.iterrows():
        rows.append({
            "line": _clean(rec.get("Line")),
            "highest_loading": _format_pct(rec.get("Highest Loading")),
            "batch": _format_num(rec.get("Batch Run no"), 0),
            "wind_speed": _format_num(rec.get("Wind Speed"), 0),
            "wind_direction": _format_num(rec.get("Wind Screen Direction"), 0),
            "water_level": _format_num(rec.get("Water Level"), 2),
            "draft": _format_num(rec.get("Draft"), 1),
        })
    return rows


def build_summary(input_data: InputReportData, output_data: OutputReportData) -> dict[str, Any]:
    """Compute KPI summary for section 4."""
    max_pct_mbl = 0.0
    max_pct_line = ""
    if not output_data.greatest_lines.empty and "Highest Loading" in output_data.greatest_lines.columns:
        series = pd.to_numeric(output_data.greatest_lines["Highest Loading"], errors="coerce")
        if series.notna().any():
            max_pct_mbl = float(series.max())
            idx = series.idxmax()
            max_pct_line = _clean(output_data.greatest_lines.loc[idx, "Line"])

    max_bollard_ton = 0.0
    max_bollard_id = ""
    if not output_data.greatest_bollards.empty and "Force" in output_data.greatest_bollards.columns:
        series = pd.to_numeric(output_data.greatest_bollards["Force"], errors="coerce")
        if series.notna().any():
            max_bollard_ton = float(series.max())
            idx = series.idxmax()
            max_bollard_id = _clean(output_data.greatest_bollards.loc[idx, "Bollard"])

    status = "OK" if max_pct_mbl <= MBL_LIMIT else "NG"
    mbl_limit_pct = f"{MBL_LIMIT * 100:.0f}%"

    return {
        "case_count": output_data.case_count,
        "vessel_name": input_data.vessel_id or input_data.vessel_name,
        "max_pct_mbl": _format_pct(max_pct_mbl),
        "max_pct_mbl_raw": max_pct_mbl,
        "max_pct_line": max_pct_line,
        "max_bollard_ton": _format_num(max_bollard_ton, 1),
        "max_bollard_id": max_bollard_id,
        "mbl_ton": _format_num(input_data.mbl_ton, 1),
        "mbl_limit_pct": mbl_limit_pct,
        "status": status,
        "status_text": (
            f"Mooring line utilization is within the {mbl_limit_pct} MBL limit."
            if status == "OK"
            else f"Mooring line utilization exceeds the {mbl_limit_pct} MBL limit — review required."
        ),
    }


def build_report_context(
    input_data: InputReportData,
    output_data: OutputReportData,
    metadata: ReportMetadata | None = None,
) -> dict[str, Any]:
    """Merge parsed data into a docxtpl render context."""
    meta = metadata or ReportMetadata()
    if not meta.report_date:
        meta.report_date = date.today().strftime("%d-%b-%Y")

    summary = build_summary(input_data, output_data)
    case_count = max(output_data.case_count, len(input_data.scenarios))

    case_headers = [f"Case {i}" for i in range(1, case_count + 1)]

    return {
        "project_name": meta.project_name,
        "reference": meta.reference,
        "report_date": meta.report_date,
        "client": meta.client,
        "author": meta.author,
        "vessel_name": input_data.vessel_id or input_data.vessel_name,
        "berth_rows": input_data.berth_rows,
        "vessel_rows": input_data.vessel_rows,
        "environment_rows": input_data.environment_rows,
        "scenarios": input_data.scenarios,
        "greatest_lines": build_greatest_line_rows(output_data.greatest_lines),
        "bollard_matrix": output_data.bollard_matrix,
        "case_headers": case_headers,
        "case_count": case_count,
        "summary": summary,
        # Flatten summary for simple {{ }} tags in template
        **{f"summary_{k}": v for k, v in summary.items()},
    }


def _heading(doc: Document, en: str, vi: str = "") -> None:
    p = doc.add_paragraph()
    run = p.add_run(en)
    run.bold = True
    run.font.size = Pt(14)
    if vi:
        p2 = doc.add_paragraph(vi)
        if p2.runs:
            p2.runs[0].italic = True
            p2.runs[0].font.size = Pt(11)


def _subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not rows:
        doc.add_paragraph("(No data)")
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val


def _render_document(context: dict[str, Any]) -> Document:
    """Build the mooring report Word document from render context."""
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("MOORING ANALYSIS REPORT\nBÁO CÁO PHÂN TÍCH NEO TÀU")
    t_run.bold = True
    t_run.font.size = Pt(18)

    for label, key in [
        ("Project / Dự án", "project_name"),
        ("Client / Chủ đầu tư", "client"),
        ("Vessel / Tàu", "vessel_name"),
        ("Reference / Số hiệu", "reference"),
        ("Date / Ngày", "report_date"),
        ("Author / Người lập", "author"),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(context.get(key, "")))

    doc.add_page_break()

    _heading(doc, "1. Input Data", "1. Số liệu đầu vào")

    _subheading(doc, "1.1 Berth Parameters / Thông số bến")
    _add_table(
        doc,
        ["Parameter", "Unit", "Value"],
        [[r["parameter"], r["unit"], r["value"]] for r in context.get("berth_rows", [])],
    )

    _subheading(doc, "1.2 Vessel Parameters / Thông số tàu")
    _add_table(
        doc,
        ["Parameter", "Unit", "Value"],
        [[r["parameter"], r["unit"], r["value"]] for r in context.get("vessel_rows", [])],
    )

    _subheading(doc, "1.3 Environmental Conditions / Điều kiện môi trường")
    _add_table(
        doc,
        ["Parameter", "Unit", "Value"],
        [[r["parameter"], r["unit"], r["value"]] for r in context.get("environment_rows", [])],
    )

    doc.add_page_break()

    _heading(doc, "2. Design Scenarios", "2. Kịch bản thiết kế")
    _add_table(
        doc,
        ["Case", "Name", "Vessel", "Loading", "Draft (m)", "Water Level (m)", "Wind Dir.", "Current Dir."],
        [
            [
                str(s["no"]), s["name"], s["vessel"], s["loading"],
                s["draft"], s["water_level"], s["wind_direction"], s["current_direction"],
            ]
            for s in context.get("scenarios", [])
        ],
    )

    doc.add_page_break()

    _heading(doc, "3. Analysis Results", "3. Kết quả phân tích")

    _subheading(doc, "3.1 Greatest Line Tensions / Lực dây neo lớn nhất (% MBL)")
    _add_table(
        doc,
        ["Line", "Highest Loading", "Batch", "Wind (kn)", "Wind Dir.", "Water Level", "Draft (m)"],
        [
            [
                r["line"], r["highest_loading"], r["batch"], r["wind_speed"],
                r["wind_direction"], r["water_level"], r["draft"],
            ]
            for r in context.get("greatest_lines", [])
        ],
    )

    _subheading(doc, "3.2 Mooring Point Loads / Tải trọng bollard (ton)")
    case_count = context.get("case_count", 0)
    bollard_headers = ["Mooring Point"] + [f"Case {i}" for i in range(1, case_count + 1)]
    bollard_rows = []
    for row in context.get("bollard_matrix", []):
        bollard_rows.append(
            [row["point"]] + [row.get(f"case_{i}", "-") for i in range(1, case_count + 1)]
        )
    _add_table(doc, bollard_headers, bollard_rows)

    doc.add_page_break()

    _heading(doc, "4. Summary", "4. Tóm tắt")
    summary = context.get("summary", {})
    summary_items = [
        ("Vessel analysed", summary.get("vessel_name", "")),
        ("Number of cases", str(summary.get("case_count", ""))),
        ("MBL (ton)", summary.get("mbl_ton", "")),
        ("Maximum line utilization", f"{summary.get('max_pct_mbl', '')} (Line {summary.get('max_pct_line', '')})"),
        ("Maximum bollard force", f"{summary.get('max_bollard_ton', '')} ton (Bollard {summary.get('max_bollard_id', '')})"),
        ("Allowable limit (OCIMF)", f"{summary.get('mbl_limit_pct', '')} MBL"),
        ("Overall status", summary.get("status", "")),
    ]
    for label, val in summary_items:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(val))

    doc.add_paragraph()
    doc.add_paragraph(str(summary.get("status_text", "")))

    return doc


def generate_report(
    input_path: Path,
    output_path: Path,
    template_path: Path,
    save_path: Path,
    metadata: ReportMetadata | None = None,
) -> Path:
    """Generate a mooring report .docx from input/output Excel files."""
    errors = validate_input_files(input_path, output_path)
    if errors:
        raise ValueError("\n".join(errors))

    input_data = read_input_workbook(input_path)
    output_data = read_output_workbook(output_path)
    context = build_report_context(input_data, output_data, metadata)

    save_path.parent.mkdir(parents=True, exist_ok=True)
    doc = _render_document(context)
    doc.save(str(save_path))
    return save_path
